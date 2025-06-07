"""Story Downloader GUI application."""

import os
import sys
from dataclasses import dataclass
from typing import Optional

from PyQt6 import QtWidgets, QtCore
import requests
from bs4 import BeautifulSoup
from docx import Document


@dataclass
class DownloadConfig:
    base_url: str
    slug: str
    start_chapter: int = 1
    css_selector: str = "p"
    user_agent: str = "Mozilla/5.0"
    timeout: int = 10
    retries: int = 3


class DownloadThread(QtCore.QThread):
    progress = QtCore.pyqtSignal(int, int)
    log = QtCore.pyqtSignal(str)
    finished = QtCore.pyqtSignal()

    def __init__(self, cfg: DownloadConfig, parent: Optional[QtCore.QObject] = None):
        super().__init__(parent)
        self.cfg = cfg
        self._stop = False

    def run(self) -> None:
        session = requests.Session()
        session.headers.update({"User-Agent": self.cfg.user_agent})

        chapter = self.cfg.start_chapter
        while not self._stop:
            url = f"{self.cfg.base_url}/{self.cfg.slug}/{chapter}"
            for attempt in range(1, self.cfg.retries + 1):
                try:
                    r = session.get(url, timeout=self.cfg.timeout)
                    if r.status_code == 404:
                        self.log.emit("Reached end of story")
                        self.finished.emit()
                        return
                    r.raise_for_status()
                    self._save_chapter(chapter, r.text)
                    self.progress.emit(chapter, 0)
                    break
                except Exception as exc:
                    self.log.emit(f"Error on chapter {chapter} attempt {attempt}: {exc}")
                    if attempt == self.cfg.retries:
                        self.finished.emit()
                        return
            chapter += 1

    def stop(self) -> None:
        self._stop = True

    def _save_chapter(self, chapter: int, html: str) -> None:
        soup = BeautifulSoup(html, "html.parser")
        title = soup.find("h1") or soup.find("title")
        title_text = title.get_text(strip=True) if title else f"Chapter {chapter}"
        doc = Document()
        doc.add_heading(title_text, level=1)
        for p in soup.select(self.cfg.css_selector):
            doc.add_paragraph(p.get_text(strip=True))
        os.makedirs(self.cfg.slug, exist_ok=True)
        doc.save(os.path.join(self.cfg.slug, f"chapter_{chapter:03}.docx"))
        self.log.emit(f"Saved chapter {chapter}")


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Story Downloader")

        self.base_url = QtWidgets.QLineEdit("https://mywebsite.com")
        self.slug = QtWidgets.QLineEdit()
        self.start_btn = QtWidgets.QPushButton("Start Download")
        self.stop_btn = QtWidgets.QPushButton("Stop")
        self.progress_bar = QtWidgets.QProgressBar()
        self.log_box = QtWidgets.QTextEdit(readOnly=True)

        layout = QtWidgets.QVBoxLayout()
        layout.addWidget(QtWidgets.QLabel("Base URL"))
        layout.addWidget(self.base_url)
        layout.addWidget(QtWidgets.QLabel("Story Slug"))
        layout.addWidget(self.slug)
        layout.addWidget(self.start_btn)
        layout.addWidget(self.stop_btn)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.log_box)
        central = QtWidgets.QWidget()
        central.setLayout(layout)
        self.setCentralWidget(central)

        self.thread: Optional[DownloadThread] = None
        self.start_btn.clicked.connect(self.start_download)
        self.stop_btn.clicked.connect(self.stop_download)

    def start_download(self):
        cfg = DownloadConfig(
            base_url=self.base_url.text().strip(),
            slug=self.slug.text().strip() or "story",
        )
        self.thread = DownloadThread(cfg)
        self.thread.progress.connect(self.update_progress)
        self.thread.log.connect(self.append_log)
        self.thread.finished.connect(self.finish)
        self.thread.start()
        self.start_btn.setEnabled(False)

    def stop_download(self):
        if self.thread and self.thread.isRunning():
            self.thread.stop()
            self.thread.wait()
            self.append_log("Download stopped")
            self.start_btn.setEnabled(True)

    def update_progress(self, chapter: int, _max: int):
        self.progress_bar.setValue(chapter)

    def append_log(self, text: str):
        self.log_box.append(text)

    def finish(self):
        QtWidgets.QMessageBox.information(self, "Done", "Download complete")
        self.start_btn.setEnabled(True)


def main() -> None:
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
